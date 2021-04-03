from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture(
    "pp.png", 
    width = Inches(2.0)
)

# name, phone number and email details
speak("Hi, I would be making your CV so please start by entering your name")
name = input("What is your name? ")

speak("Hello " + name + " , can you give me your phone number please?")
phone_number = input("What is your phone number? ")

speak("Thank you, now please enter your email")
email = input("What is your email? ")
speak("Thanks")

document.add_paragraph(
    name + " | " + phone_number + " | " + email)

# about me
document.add_heading("About me")
speak("Can I get more information by you telling me about yourself please?")
document.add_paragraph(
    input("Tell about yourself ")
)
speak("Thank you")

# work experience
speak("Let's see your work experience, start by entering the name of the company at which you have worked")
document.add_heading("Work Experience")
p = document.add_paragraph()

company = input("Enter the name of the company ")
speak("Ok, at what year did you started working at " + company + "?")
from_date = input("From (year) ")
speak("At what year did you stopped working for " + company + "?")
to_date = input("To (year) ")

p.add_run(company + " ").bold = True
p.add_run(from_date + "-" + to_date + "\n").italic = True

speak("Can you please enter your position and write how did you performed at the company?")
experience_details = input(
    "Describe your experience at " + company + " ")
p.add_run(experience_details)
speak("Thank you for providing the required information")

# more experiences
while True:
    speak("Have you worked in another company besides the one you mentioned?")
    has_more_experiences = input(
        "Do you have more experience? Yes or No - ")
    if has_more_experiences.lower() == "yes":
        p = document.add_paragraph()

        speak("What was the company's name")
        company = input("Enter the name of the company ")
        speak("Ok, at what year did you started working at " + company + "?")
        from_date = input("From (year) ")
        speak("At what year did you stopped working for " + company + "?")
        to_date = input("To (year) ")
        speak("OK")

        p.add_run(company + " ").bold = True
        p.add_run(from_date + "-" + to_date + "\n").italic = True

        speak("Can you please enter your position and write how did you performed at the company")
        experience_details = input(
            "Describe your experience at " + company + " ")
        p.add_run(experience_details)
    else:
        break

# skills
document.add_heading("Skills")
speak("Now let's talk about the skills that you have, please enter a skill")
skill = input("Enter a skill you have ")
p = document.add_paragraph(skill)
p.style = "List Bullet"

# more skills
while True:
    speak("Do you have more skills that you want to share?")
    has_more_skills = input("Do you have more skills? Yes or No - ")
    if has_more_skills.lower() == "yes":
        speak("Please write about your other skill")
        skill = input("Enter your other skill ")
        speak("Thank you")
        p = document.add_paragraph(skill)
        p.style = "List Bullet"
    else:
        break

speak("Your CV is completely ready, thank you for using Ari's services")

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "This CV was generated thanks to the mastermind of Ari"

document.save("cv.docx")